from reportlab.platypus import Paragraph
from app.constants.resume_constants import COMPANY_HEADING_PARAGRAPH_STYLE, COMPANY_DURATION_PARAGRAPH_STYLE, COMPANY_TITLE_PARAGRAPH_STYLE, COMPANY_LOCATION_PARAGRAPH_STYLE, JOB_DETAILS_PARAGRAPH_STYLE
from reportlab.lib.styles import ParagraphStyle
from app.constants import GARAMOND_REGULAR, GARAMOND_SEMIBOLD
from reportlab.lib.enums import TA_JUSTIFY
from docx.shared import Pt

# Style for consulting skill headers (bold inline)
SKILL_HEADER_BULLET_STYLE = ParagraphStyle(
    'skill_header_bullet',
    leftIndent=12,
    fontName=GARAMOND_REGULAR,
    fontSize=11,
    leading=12,
    alignment=TA_JUSTIFY
)


class ConsultingExperience:
    """Experience class for consulting resume format with skill headers."""

    def __init__(self, company='', title='', location='', start_date='', end_date='', description=[]) -> None:
        self.company = company
        self.title = title
        self.location = location
        self.start_date = start_date
        self.end_date = end_date
        self.description = description  # List of dicts with 'skillHeader' and 'bullet'

    def set_company(self, company: str) -> None:
        self.company = company

    def set_title(self, title: str) -> None:
        self.title = title

    def set_location(self, location: str) -> None:
        self.location = location

    def set_start_date(self, start_date: str) -> None:
        self.start_date = start_date

    def set_end_date(self, end_date: str) -> None:
        self.end_date = end_date

    def set_description(self, description: list) -> None:
        self.description = description

    def append_description(self, item: dict) -> None:
        self.description.append(item)

    def __str__(self) -> str:
        return f"{{company: {self.company}, title: {self.title}, location: {self.location}, start_date: {self.start_date}, end_date: {self.end_date}, description: {self.description}}}"

    def get_table_element(self, running_row_index: list, table_styles: list) -> list:
        experience_table = []

        # Parse company field to extract company name and location (format: "Company | Location")
        company_name = self.company
        location = ''
        if ' | ' in self.company:
            parts = self.company.split(' | ', 1)
            company_name = parts[0].strip()
            location = parts[1].strip()

        experience_table.append([
            Paragraph(company_name, COMPANY_HEADING_PARAGRAPH_STYLE),
            Paragraph(location, COMPANY_DURATION_PARAGRAPH_STYLE)
        ])
        table_styles.append(('TOPPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 2))
        table_styles.append(('BOTTOMPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 0))
        running_row_index[0] += 1

        experience_table.append([
            Paragraph(self.title, COMPANY_TITLE_PARAGRAPH_STYLE),
            Paragraph(f"{self.start_date} - {self.end_date}", COMPANY_DURATION_PARAGRAPH_STYLE)
        ])
        table_styles.append(('TOPPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 0))
        table_styles.append(('BOTTOMPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 0))
        running_row_index[0] += 1

        for item in self.description:
            # Handle consulting format with skillHeader and bullet
            if isinstance(item, dict) and 'skillHeader' in item:
                skill_header = item.get('skillHeader', '')
                bullet_text = item.get('bullet', '')
                # Format: "SkillHeader: Bullet text" with skill header in bold using explicit font
                formatted_text = f'<font name="{GARAMOND_SEMIBOLD}">{skill_header}:</font> {bullet_text}'
                experience_table.append([
                    Paragraph(formatted_text, bulletText='•', style=SKILL_HEADER_BULLET_STYLE), ''
                ])
            elif isinstance(item, str):
                # Fallback for plain string descriptions
                experience_table.append([
                    Paragraph(item, bulletText='•', style=JOB_DETAILS_PARAGRAPH_STYLE), ''
                ])
            else:
                continue

            table_styles.append(('TOPPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 0))
            table_styles.append(('BOTTOMPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 0))
            table_styles.append(('SPAN', (0, running_row_index[0]), (1, running_row_index[0])))
            running_row_index[0] += 1

        return experience_table

    def get_docx_content(self, doc):
        """Add experience content to DOCX document"""
        # Parse company field to extract company name and location (format: "Company | Location")
        company_name = self.company
        location = ''
        if ' | ' in self.company:
            parts = self.company.split(' | ', 1)
            company_name = parts[0].strip()
            location = parts[1].strip()

        # Company name and location on same line
        company_paragraph = doc.add_paragraph()
        company_run = company_paragraph.add_run(company_name)
        company_run.font.size = Pt(11)
        company_run.font.bold = True
        company_run.font.name = 'Calibri'

        # Add location on the same line, right-aligned (bold)
        if location:
            location_run = company_paragraph.add_run(f"\t{location}")
            location_run.font.size = Pt(11)
            location_run.font.bold = True
            location_run.font.name = 'Calibri'

        # Job title and period
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(self.title)
        title_run.font.size = Pt(11)
        title_run.font.italic = True
        title_run.font.name = 'Calibri'

        # Add dates on the same line, right-aligned
        dates_run = title_paragraph.add_run(f"\t{self.start_date} - {self.end_date}")
        dates_run.font.size = Pt(11)
        dates_run.font.name = 'Calibri'

        # Description bullets with skill headers
        for item in self.description:
            if isinstance(item, dict) and 'skillHeader' in item:
                skill_header = item.get('skillHeader', '')
                bullet_text = item.get('bullet', '')

                desc_paragraph = doc.add_paragraph()
                # Add bullet
                bullet_run = desc_paragraph.add_run("• ")
                bullet_run.font.size = Pt(11)
                bullet_run.font.name = 'Calibri'

                # Add skill header (bold)
                header_run = desc_paragraph.add_run(f"{skill_header}: ")
                header_run.font.size = Pt(11)
                header_run.font.bold = True
                header_run.font.name = 'Calibri'

                # Add bullet text
                text_run = desc_paragraph.add_run(bullet_text)
                text_run.font.size = Pt(11)
                text_run.font.name = 'Calibri'
            elif isinstance(item, str) and item.strip():
                desc_paragraph = doc.add_paragraph()
                desc_run = desc_paragraph.add_run(f"• {item}")
                desc_run.font.size = Pt(11)
                desc_run.font.name = 'Calibri'

        # Add space after experience
        doc.add_paragraph()
