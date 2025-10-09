from reportlab.platypus import Paragraph
from app.constants.resume_constants import COMPANY_HEADING_PARAGRAPH_STYLE, COMPANY_DURATION_PARAGRAPH_STYLE, COMPANY_TITLE_PARAGRAPH_STYLE, COMPANY_LOCATION_PARAGRAPH_STYLE, JOB_DETAILS_PARAGRAPH_STYLE
from docx.shared import Pt

class Education:
    def __init__(self, institution='', course='', location='', start_date='', end_date='') -> None:
        self.institution = institution
        self.course = course
        self.location = location
        self.start_date = start_date
        self.end_date = end_date
        
    def set_institution(self, institution : str) -> None:
        self.institution = institution
        
    def set_course(self, course : str) -> None:
        self.course = course
        
    def set_location(self, location : str) -> None:
        self.location = location
        
    def set_start_date(self, start_date : str) -> None:
        self.start_date = start_date
        
    def set_end_date(self, end_date : str) -> None:
        self.end_date = end_date
        
    def get_table_element(self, running_row_index : list, table_styles : list) -> list:
        education_table = []
        education_table.append([
            Paragraph(self.institution, COMPANY_HEADING_PARAGRAPH_STYLE),
            Paragraph(f"{self.start_date} - {self.end_date}", COMPANY_DURATION_PARAGRAPH_STYLE)
        ])
        table_styles.append(('TOPPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 5))
        running_row_index[0] += 1
        
        education_table.append([
            Paragraph(self.course, COMPANY_TITLE_PARAGRAPH_STYLE),
            Paragraph(self.location, COMPANY_LOCATION_PARAGRAPH_STYLE)
        ])
        table_styles.append(('TOPPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 1))
        running_row_index[0] += 1
        
        return education_table
    
    def get_docx_content(self, doc):
        """Add education content to DOCX document"""
        # Institution and dates
        institution_paragraph = doc.add_paragraph()
        institution_run = institution_paragraph.add_run(self.institution)
        institution_run.font.size = Pt(11)
        institution_run.font.bold = True
        institution_run.font.name = 'Calibri'
        
        if self.start_date or self.end_date:
            dates_run = institution_paragraph.add_run(f"\t{self.start_date} - {self.end_date}")
            dates_run.font.size = Pt(11)
            dates_run.font.name = 'Calibri'
        
        # Course/Degree and location
        course_paragraph = doc.add_paragraph()
        course_run = course_paragraph.add_run(self.course)
        course_run.font.size = Pt(11)
        course_run.font.italic = True
        course_run.font.name = 'Calibri'
        
        if self.location:
            location_run = course_paragraph.add_run(f"\t{self.location}")
            location_run.font.size = Pt(11)
            location_run.font.name = 'Calibri'
        
        # Add space after education
        doc.add_paragraph()