from reportlab.platypus import Paragraph
from app.constants.resume_constants import JOB_DETAILS_PARAGRAPH_STYLE
from docx.shared import Pt

class Skill:
    def __init__(self, title='', elements=[]) -> None:
        self.title = title
        self.elements = elements
        
    def set_title(self, title : str) -> None:
        self.title = title
        
    def set_elements(self, elements : list) -> None:
        self.elements = elements
    
    def append_element(self, element : str) -> None:
        self.elements.append(element)
        
    def get_table_element(self, running_row_index : list, table_styles : list) -> list:
        table = []
        # Fix the syntax error in joining elements
        elements_string = ", ".join(word for word in self.elements if word)
        table.append([
            Paragraph(f"<font face='Garamond_Semibold'>{self.title}:</font> {elements_string}", bulletText='•', style=JOB_DETAILS_PARAGRAPH_STYLE)
        ])
        table_styles.append(('TOPPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 1))
        table_styles.append(('BOTTOMPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 0))
        table_styles.append(('SPAN', (0, running_row_index[0]), (1, running_row_index[0])))
        running_row_index[0] += 1
        return table
    
    def get_docx_content(self, doc):
        """Add skill content to DOCX document"""
        skill_paragraph = doc.add_paragraph()
        
        # Add skill category title in bold
        title_run = skill_paragraph.add_run(f"{self.title}: ")
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.name = 'Calibri'
        
        # Add skills list
        elements_string = ", ".join(word for word in self.elements if word)
        skills_run = skill_paragraph.add_run(elements_string)
        skills_run.font.size = Pt(11)
        skills_run.font.name = 'Calibri'