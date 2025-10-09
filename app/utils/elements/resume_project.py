from reportlab.platypus import Paragraph
from app.constants.resume_constants import JOB_DETAILS_PARAGRAPH_STYLE
from docx.shared import Pt


class Project:
    def __init__(self, title='', description='', link='') -> None:
        self.title = title
        self.description = description
        self.link = link
        
    def set_title(self, title : str) -> None:
        self.title = title
        
    def set_description(self, description : str) -> None:
        self.description = description
        
    def set_link(self, link : str) -> None:
        self.link = link
        
    def get_table_element(self, running_row_index : list, table_styles : list) -> list:
        table = []
        table.append([
            Paragraph(f"<font face='Garamond_Semibold'>{self.title}: </font>{self.description} {self.link}", bulletText='•', style=JOB_DETAILS_PARAGRAPH_STYLE),
        ])
        table_styles.append(('TOPPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 1))
        table_styles.append(('BOTTOMPADDING', (0, running_row_index[0]), (1, running_row_index[0]), 0))
        table_styles.append(('SPAN', (0, running_row_index[0]), (1, running_row_index[0])))
        running_row_index[0] += 1
        return table
    
    def get_docx_content(self, doc):
        """Add project content to DOCX document"""
        project_paragraph = doc.add_paragraph()
        
        # Add project title in bold
        title_run = project_paragraph.add_run(f"{self.title}: ")
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.name = 'Calibri'
        
        # Add description
        desc_run = project_paragraph.add_run(self.description)
        desc_run.font.size = Pt(11)
        desc_run.font.name = 'Calibri'
        
        # Add link if available
        if self.link:
            link_run = project_paragraph.add_run(f" {self.link}")
            link_run.font.size = Pt(11)
            link_run.font.name = 'Calibri'