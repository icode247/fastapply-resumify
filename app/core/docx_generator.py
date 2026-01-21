"""
DOCX resume generation functionality for ATS compatibility.
"""
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from app.constants.resume_constants import (
    ATS_RESUME_ELEMENTS_ORDER, JAKE_SECTION_ORDER_NEW_GRAD, JAKE_SECTION_ORDER_EXPERIENCED,
    HARVARD_SECTION_ORDER
)
from app.utils.helpers import get_education_element, get_experience_element, get_consulting_experience_element, get_project_element, get_skills_element, get_achievements_element
from app.utils.sections.resume_section import Section


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph in DOCX

    Args:
        paragraph: The paragraph to add the hyperlink to
        text: The text to display for the link
        url: The URL to link to

    Returns:
        The hyperlink element
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = OxmlElement('w:r')

    # Set the run's text
    rPr = OxmlElement('w:rPr')

    # Set the run's style to 'Hyperlink'
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink


def generate_resume_docx(author, resume_data):
    """
    Generate a DOCX resume optimized for ATS systems
    
    Args:
        author (str): Name of the person
        resume_data (dict): Resume data containing education, experience, projects, skills and contact info
        
    Returns:
        bytes: The generated DOCX content as bytes
    """
    # Create new document
    doc = Document()
    
    # Set up document margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Use provided author or get from resume_data
    if not author and 'name' in resume_data:
        author = resume_data.get('name', '')
    
    # Extract contact information from nested contact object
    contact = resume_data.get('contact', {})
    email = contact.get('email', '')
    phone = contact.get('phone', '')
    location = contact.get('location', '')
    linkedin = contact.get('linkedin', '')
    github = contact.get('github', '')
    portfolio = contact.get('portfolio', '')
    job_title = resume_data.get('title', '')
    
    # Add name (header)
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(author)
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add job title if available
    if job_title:
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(job_title)
        title_run.font.size = Pt(12)
        title_run.font.name = 'Calibri'
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add contact information
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build contact information with hyperlinks
    first_item = True

    if email:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(email)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if phone:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(phone)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if location:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(location)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if github:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        add_hyperlink(contact_paragraph, "Github", github)
        first_item = False

    if linkedin:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        add_hyperlink(contact_paragraph, "Linkedin", linkedin)
        first_item = False

    if portfolio:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(portfolio)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False
    
    # Add some space after contact info
    doc.add_paragraph()

    # Add Professional Summary (use the one from resume_data if available)
    summary_text = resume_data.get('summary', '')
    if summary_text:
        add_section_header(doc, "PROFESSIONAL SUMMARY")
        summary_paragraph = doc.add_paragraph()
        summary_run = summary_paragraph.add_run(summary_text)
        summary_run.font.size = Pt(11)
        summary_run.font.name = 'Calibri'
        doc.add_paragraph()
    
    # Process resume sections in ATS-optimized order
    processed_resume_data = process_resume_sections(resume_data)
    
    # Add sections in proper ATS order
    for element in ATS_RESUME_ELEMENTS_ORDER:
        if element in processed_resume_data:
            section = processed_resume_data[element]
            add_resume_section_to_doc(doc, section)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def process_resume_sections(resume_data):
    """Process resume data into sections"""
    processed_resume_data = {}
    
    # Process experience data
    experience_elements = []
    if 'experience' in resume_data and resume_data['experience']:
        for element in resume_data['experience']:
            experience_elements.append(get_experience_element(element))
        if experience_elements:  # Only add section if there are elements
            processed_resume_data['experience'] = Section('PROFESSIONAL EXPERIENCE', experience_elements)
    
    # Process skills data
    skill_elements = []
    if 'skills' in resume_data:
        skills_data = resume_data['skills']
        
        if isinstance(skills_data, dict):
            # Process technical skills
            if 'technical' in skills_data and skills_data['technical']:
                skill_elements.append(get_skills_element('Technical Skills', skills_data['technical']))

            # Process programming languages
            if 'languages' in skills_data and skills_data['languages']:
                skill_elements.append(get_skills_element('Programming Languages', skills_data['languages']))

            # Process frameworks
            if 'frameworks' in skills_data and skills_data['frameworks']:
                skill_elements.append(get_skills_element('Frameworks & Libraries', skills_data['frameworks']))
            elif 'frameworks/libraries' in skills_data and skills_data['frameworks/libraries']:
                skill_elements.append(get_skills_element('Frameworks & Libraries', skills_data['frameworks/libraries']))

            # Process tools
            if 'tools' in skills_data and skills_data['tools']:
                skill_elements.append(get_skills_element('Tools', skills_data['tools']))

            # Process technologies (legacy field)
            if 'technologies' in skills_data and skills_data['technologies']:
                skill_elements.append(get_skills_element('Technologies', skills_data['technologies']))

            # Process methodologies
            if 'methodologies' in skills_data and skills_data['methodologies']:
                skill_elements.append(get_skills_element('Methodologies', skills_data['methodologies']))

            # Process soft skills
            if 'soft_skills' in skills_data and skills_data['soft_skills']:
                skill_elements.append(get_skills_element('Soft Skills', skills_data['soft_skills']))

            # Process others (legacy field)
            if 'others' in skills_data and skills_data['others']:
                skill_elements.append(get_skills_element('Other Skills', skills_data['others']))
        
        elif isinstance(skills_data, list):
            for skill in skills_data:
                if isinstance(skill, dict) and 'title' in skill:
                    elements = skill.get('elements', [])
                    if elements:  # Only add if there are elements
                        skill_elements.append(get_skills_element(skill['title'], elements))
    
    if skill_elements:
        processed_resume_data['skills'] = Section('CORE COMPETENCIES', skill_elements)
    
    # Process education data
    education_elements = []
    if 'education' in resume_data and resume_data['education']:
        for element in resume_data['education']:
            education_elements.append(get_education_element(element))
        if education_elements:  # Only add section if there are elements
            processed_resume_data['education'] = Section('EDUCATION', education_elements)
    
    # Process projects data
    project_elements = []
    if 'projects' in resume_data and resume_data['projects'] and len(resume_data['projects']) > 0:
        for element in resume_data['projects']:
            project_elements.append(get_project_element(element))
        if project_elements:  # Only add section if there are elements
            processed_resume_data['projects'] = Section('PROJECTS', project_elements)

    # Process achievements data
    if 'achievements' in resume_data and resume_data['achievements']:
        achievements_list = resume_data['achievements']
        if isinstance(achievements_list, list) and achievements_list:
            achievement_elements = []
            achievement_elements.append(get_achievements_element(achievements_list))
            processed_resume_data['achievements'] = Section('ACHIEVEMENTS', achievement_elements)

    # Process certifications data
    if 'certifications' in resume_data and resume_data['certifications']:
        certifications = resume_data['certifications']
        if isinstance(certifications, list) and certifications:
            cert_elements = []
            cert_elements.append(get_achievements_element(certifications))
            processed_resume_data['certifications'] = Section('CERTIFICATIONS', cert_elements)

    # Process languages data (add as separate section if present)
    if 'languages' in resume_data and resume_data['languages']:
        language_list = resume_data['languages']
        if isinstance(language_list, list) and language_list:
            lang_elements = []
            lang_elements.append(get_skills_element('Languages', language_list))
            processed_resume_data['languages'] = Section('LANGUAGES', lang_elements)

    return processed_resume_data


def add_section_header(doc, header_text):
    """Add a section header with proper formatting"""
    header_paragraph = doc.add_paragraph()
    header_run = header_paragraph.add_run(header_text)
    header_run.font.size = Pt(12)
    header_run.font.bold = True
    header_run.font.name = 'Calibri'
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_resume_section_to_doc(doc, section):
    """Add a resume section to the document"""
    # Add section header
    add_section_header(doc, section.heading)
    
    # Add section content
    for element in section.elements:
        if hasattr(element, 'get_docx_content'):
            element.get_docx_content(doc)
        else:
            # Fallback for elements without docx support
            add_generic_element_to_doc(doc, element)
    
    # Add space after section
    doc.add_paragraph()


def add_generic_element_to_doc(doc, element):
    """Add a generic element to the document"""
    if hasattr(element, 'title') and element.title:
        # Add element title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(element.title)
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.name = 'Calibri'

    if hasattr(element, 'description') and element.description:
        # Add description
        if isinstance(element.description, list):
            for desc in element.description:
                desc_paragraph = doc.add_paragraph()
                desc_run = desc_paragraph.add_run(f"• {desc}")
                desc_run.font.size = Pt(11)
                desc_run.font.name = 'Calibri'
        else:
            desc_paragraph = doc.add_paragraph()
            desc_run = desc_paragraph.add_run(element.description)
            desc_run.font.size = Pt(11)
            desc_run.font.name = 'Calibri'


def process_resume_sections_with_format(resume_data, is_consulting=False):
    """Process resume data into sections with support for consulting format"""
    processed_resume_data = {}

    # Process experience data
    experience_elements = []
    if 'experience' in resume_data and resume_data['experience']:
        for element in resume_data['experience']:
            if is_consulting:
                experience_elements.append(get_consulting_experience_element(element))
            else:
                experience_elements.append(get_experience_element(element))
        if experience_elements:
            processed_resume_data['experience'] = Section('PROFESSIONAL EXPERIENCE', experience_elements)

    # Process skills data
    skill_elements = []
    if 'skills' in resume_data:
        skills_data = resume_data['skills']

        if isinstance(skills_data, dict):
            if 'technical' in skills_data and skills_data['technical']:
                skill_elements.append(get_skills_element('Technical Skills', skills_data['technical']))
            if 'languages' in skills_data and skills_data['languages']:
                skill_elements.append(get_skills_element('Programming Languages', skills_data['languages']))
            if 'frameworks' in skills_data and skills_data['frameworks']:
                skill_elements.append(get_skills_element('Frameworks & Libraries', skills_data['frameworks']))
            elif 'frameworks/libraries' in skills_data and skills_data['frameworks/libraries']:
                skill_elements.append(get_skills_element('Frameworks & Libraries', skills_data['frameworks/libraries']))
            if 'tools' in skills_data and skills_data['tools']:
                skill_elements.append(get_skills_element('Tools', skills_data['tools']))
            if 'technologies' in skills_data and skills_data['technologies']:
                skill_elements.append(get_skills_element('Technologies', skills_data['technologies']))
            if 'methodologies' in skills_data and skills_data['methodologies']:
                skill_elements.append(get_skills_element('Methodologies', skills_data['methodologies']))
            if 'soft_skills' in skills_data and skills_data['soft_skills']:
                skill_elements.append(get_skills_element('Soft Skills', skills_data['soft_skills']))
            if 'others' in skills_data and skills_data['others']:
                skill_elements.append(get_skills_element('Other Skills', skills_data['others']))

        elif isinstance(skills_data, list):
            for skill in skills_data:
                if isinstance(skill, dict) and 'title' in skill:
                    elements = skill.get('elements', [])
                    if elements:
                        skill_elements.append(get_skills_element(skill['title'], elements))

    if skill_elements:
        processed_resume_data['skills'] = Section('CORE COMPETENCIES', skill_elements)

    # Process education data
    education_elements = []
    if 'education' in resume_data and resume_data['education']:
        for element in resume_data['education']:
            education_elements.append(get_education_element(element))
        if education_elements:
            processed_resume_data['education'] = Section('EDUCATION', education_elements)

    # Process projects data
    project_elements = []
    if 'projects' in resume_data and resume_data['projects'] and len(resume_data['projects']) > 0:
        for element in resume_data['projects']:
            project_elements.append(get_project_element(element))
        if project_elements:
            processed_resume_data['projects'] = Section('PROJECTS', project_elements)

    # Process achievements data
    if 'achievements' in resume_data and resume_data['achievements']:
        achievements_list = resume_data['achievements']
        if isinstance(achievements_list, list) and achievements_list:
            achievement_elements = []
            achievement_elements.append(get_achievements_element(achievements_list))
            processed_resume_data['achievements'] = Section('ACHIEVEMENTS', achievement_elements)

    # Process certifications data
    if 'certifications' in resume_data and resume_data['certifications']:
        certifications = resume_data['certifications']
        if isinstance(certifications, list) and certifications:
            cert_elements = []
            cert_elements.append(get_achievements_element(certifications))
            processed_resume_data['certifications'] = Section('CERTIFICATIONS', cert_elements)

    # Process languages data
    if 'languages' in resume_data and resume_data['languages']:
        language_list = resume_data['languages']
        if isinstance(language_list, list) and language_list:
            lang_elements = []
            lang_elements.append(get_skills_element('Languages', language_list))
            processed_resume_data['languages'] = Section('LANGUAGES', lang_elements)

    return processed_resume_data


def generate_jake_resume_docx(author, resume_data, years_of_experience=0, is_consulting=False):
    """
    Generate a DOCX resume using Jake's template.

    Jake's template features:
    - Centered header (name, contact)
    - Compact, dense layout with 0.5" margins
    - Dynamic section order based on years_of_experience:
      - <= 2 years: Education -> Experience -> Projects -> Skills
      - > 2 years: Experience -> Skills -> Education -> Projects

    Args:
        author (str): Name of the person
        resume_data (dict): Resume data
        years_of_experience (int): Years of experience to determine section order
        is_consulting (bool): If True, use consulting experience format

    Returns:
        bytes: The generated DOCX content as bytes
    """
    # Create new document
    doc = Document()

    # Set up document margins (0.5 inch - compact for Jake's template)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Use provided author or get from resume_data
    if not author and 'name' in resume_data:
        author = resume_data.get('name', '')

    # Extract contact information
    contact = resume_data.get('contact', {})
    email = contact.get('email', '')
    phone = contact.get('phone', '')
    location = contact.get('location', '')
    linkedin = contact.get('linkedin', '')
    github = contact.get('github', '')
    portfolio = contact.get('portfolio', '')
    job_title = resume_data.get('title', '')

    # Determine section order based on years of experience
    # Less than 3 years: Education first (new grad friendly)
    # 3+ years: Experience first (professional)
    if years_of_experience < 3:
        section_order = JAKE_SECTION_ORDER_NEW_GRAD
    else:
        section_order = JAKE_SECTION_ORDER_EXPERIENCED

    # Add name (header) - centered
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(author)
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add job title if available - centered
    if job_title:
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(job_title)
        title_run.font.size = Pt(12)
        title_run.font.name = 'Calibri'
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add contact information - centered
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    first_item = True
    if email:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(email)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if phone:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(phone)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if location:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(location)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if github:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        add_hyperlink(contact_paragraph, "Github", github)
        first_item = False

    if linkedin:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        add_hyperlink(contact_paragraph, "Linkedin", linkedin)
        first_item = False

    if portfolio:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(portfolio)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    # Add optional summary section (only for 3+ years experience)
    summary_text = resume_data.get('summary', '')
    if summary_text and years_of_experience >= 3:
        add_section_header(doc, "SUMMARY")
        summary_paragraph = doc.add_paragraph()
        summary_run = summary_paragraph.add_run(summary_text)
        summary_run.font.size = Pt(11)
        summary_run.font.name = 'Calibri'

    # Process resume sections
    processed_resume_data = process_resume_sections_with_format(resume_data, is_consulting)

    # Add sections in Jake's dynamic order
    for element in section_order:
        if element in processed_resume_data:
            section = processed_resume_data[element]
            add_resume_section_to_doc(doc, section)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def generate_harvard_resume_docx(author, resume_data, years_of_experience=0, is_consulting=False):
    """
    Generate a DOCX resume using Harvard's template.

    Harvard's template features:
    - Left-aligned header
    - Summary section (optional - only shown for 3+ years experience)
    - Section order: Summary (if present) -> Experience -> Skills -> Education -> Projects
    - More whitespace with 1" margins

    Args:
        author (str): Name of the person
        resume_data (dict): Resume data
        years_of_experience (int): Years of experience (summary only shown for 3+)
        is_consulting (bool): If True, use consulting experience format

    Returns:
        bytes: The generated DOCX content as bytes
    """
    # Create new document
    doc = Document()

    # Set up document margins (1 inch - more whitespace for Harvard template)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Use provided author or get from resume_data
    if not author and 'name' in resume_data:
        author = resume_data.get('name', '')

    # Extract contact information
    contact = resume_data.get('contact', {})
    email = contact.get('email', '')
    phone = contact.get('phone', '')
    location = contact.get('location', '')
    linkedin = contact.get('linkedin', '')
    github = contact.get('github', '')
    portfolio = contact.get('portfolio', '')
    job_title = resume_data.get('title', '')

    # Add name (header) - left-aligned for Harvard
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(author)
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add job title if available - left-aligned
    if job_title:
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(job_title)
        title_run.font.size = Pt(12)
        title_run.font.name = 'Calibri'
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add contact information - left-aligned
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    first_item = True
    if email:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(email)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if phone:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(phone)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if location:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(location)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    if github:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        add_hyperlink(contact_paragraph, "Github", github)
        first_item = False

    if linkedin:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        add_hyperlink(contact_paragraph, "Linkedin", linkedin)
        first_item = False

    if portfolio:
        if not first_item:
            contact_paragraph.add_run(" | ").font.size = Pt(11)
        run = contact_paragraph.add_run(portfolio)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        first_item = False

    # Add some space after contact info
    doc.add_paragraph()

    # Add Professional Summary (only for 3+ years experience)
    summary_text = resume_data.get('summary', '')
    if summary_text and years_of_experience >= 3:
        add_section_header(doc, "PROFESSIONAL SUMMARY")
        summary_paragraph = doc.add_paragraph()
        summary_run = summary_paragraph.add_run(summary_text)
        summary_run.font.size = Pt(11)
        summary_run.font.name = 'Calibri'
        doc.add_paragraph()

    # Process resume sections
    processed_resume_data = process_resume_sections_with_format(resume_data, is_consulting)

    # Add sections in Harvard's fixed order
    for element in HARVARD_SECTION_ORDER:
        if element in processed_resume_data:
            section = processed_resume_data[element]
            add_resume_section_to_doc(doc, section)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()